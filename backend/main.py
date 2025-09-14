from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import os
from datetime import datetime
import fitz
from pdf2docx import Converter
from docx import Document
import asyncio
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv
import json
import re

load_dotenv()
app = FastAPI()

# ✅ CORS Settings for Frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Only upload and preview APIs are exposed

# ✅ File Upload Endpoint
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Track the current active document
CURRENT_DOCUMENT = None

# Serve uploads as static files
app.mount("/uploads", StaticFiles(directory=UPLOAD_FOLDER), name="uploads")

async def _convert_pdf_to_docx(pdf_path: str, docx_path: str):
    loop = asyncio.get_running_loop()
    def _blocking_convert():
        cv = Converter(pdf_path)
        try:
            cv.convert(docx_path)
        finally:
            cv.close()
    # Run the blocking conversion in a thread to avoid blocking the event loop
    await loop.run_in_executor(None, _blocking_convert)

def _modify_document(doc_path: str, modification_prompt: str, groq_client):
    """Modify the document based on the AI prompt"""
    try:
        # Read the current document
        doc = Document(doc_path)
        
        # Extract current content for context
        current_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                current_content.append(paragraph.text)
        
        # Create a comprehensive prompt for the AI
        system_prompt = """You are an expert resume editor. You will receive a document and modification instructions. 
        You must modify the document according to the instructions while maintaining the original structure and formatting.
        
        IMPORTANT: You must return ONLY the modified content in JSON format with this exact structure:
        {
            "paragraphs": ["modified paragraph 1", "modified paragraph 2", ...],
            "tables": [
                {
                    "rows": [
                        ["cell1", "cell2", ...],
                        ["cell1", "cell2", ...]
                    ]
                }
            ]
        }
        
        If there are no tables, return empty array for tables. Preserve the logical flow and professional tone."""
        
        user_prompt = f"""Current document content:
{chr(10).join(current_content)}

Modification request: {modification_prompt}

Please modify the document according to the request and return the updated content in the specified JSON format."""
        
        # Get AI response
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            model="llama-3.3-70b-versatile",
        )
        
        ai_response = chat_completion.choices[0].message.content
        
        # Parse AI response
        try:
            # Extract JSON from the response (sometimes AI adds extra text)
            json_match = re.search(r'\{.*\}', ai_response, re.DOTALL)
            if json_match:
                modifications = json.loads(json_match.group())
            else:
                raise ValueError("No JSON found in AI response")
            
            # Clear existing content
            doc._element.clear()
            
            # Add modified paragraphs
            for para_text in modifications.get("paragraphs", []):
                if para_text.strip():
                    doc.add_paragraph(para_text)
            
            # Add modified tables
            for table_data in modifications.get("tables", []):
                if table_data.get("rows"):
                    table = doc.add_table(rows=len(table_data["rows"]), cols=len(table_data["rows"][0]))
                    for i, row_data in enumerate(table_data["rows"]):
                        for j, cell_data in enumerate(row_data):
                            if i < len(table.rows) and j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = str(cell_data)
            
            # Save the modified document
            doc.save(doc_path)
            return True
            
        except (json.JSONDecodeError, KeyError, ValueError) as e:
            print(f"Error parsing AI response: {e}")
            print(f"AI Response: {ai_response}")
            return False
            
    except Exception as e:
        print(f"Error modifying document: {e}")
        return False

class ChatRequest(BaseModel):
    message: str

@app.post("/chat")
async def chat_endpoint(chat: ChatRequest):
    groq_key = os.getenv("GROQ_API_KEY")
    if not groq_key:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY not configured on server")
    
    try:
        client = Groq(api_key=groq_key)
        
        # Check if this is a document modification request
        modification_keywords = [
            "modify", "edit", "change", "update", "improve", "rewrite", 
            "resume", "document", "cv", "curriculum vitae"
        ]
        
        is_modification_request = any(keyword in chat.message.lower() for keyword in modification_keywords)
        
        if is_modification_request and CURRENT_DOCUMENT:
            # This is a document modification request
            doc_path = os.path.join(UPLOAD_FOLDER, CURRENT_DOCUMENT)
            
            if os.path.exists(doc_path):
                # Modify the document
                success = await asyncio.get_running_loop().run_in_executor(
                    None, _modify_document, doc_path, chat.message, client
                )
                
                if success:
                    # After successful modification, get the updated preview content
                    try:
                        # Read the updated document content
                        doc = Document(doc_path)
                        updated_content = []
                        
                        # Extract paragraphs
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                updated_content.append(paragraph.text)
                        
                        # Extract table content
                        table_content = []
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = [cell.text for cell in row.cells]
                                table_content.append(" | ".join(row_text))
                        
                        # Combine all content
                        full_content = "\n".join(updated_content)
                        if table_content:
                            full_content += "\n\nTables:\n" + "\n".join(table_content)
                        
                        return {
                            "response": f"Document '{CURRENT_DOCUMENT}' has been successfully modified according to your request. Here's the updated content:",
                            "document_modified": True,
                            "filename": CURRENT_DOCUMENT,
                            "updated_content": full_content[:10000],  # Limit content length
                            "content_type": "docx"
                        }
                        
                    except Exception as e:
                        # If preview generation fails, still return success but without content
                        return {
                            "response": f"Document '{CURRENT_DOCUMENT}' has been successfully modified according to your request. You can now preview the updated document.",
                            "document_modified": True,
                            "filename": CURRENT_DOCUMENT,
                            "updated_content": None,
                            "content_type": "docx"
                        }
                else:
                    return {
                        "response": "I encountered an error while modifying the document. Please try again with a clearer instruction.",
                        "document_modified": False
                    }
            else:
                return {
                    "response": "No active document found to modify. Please upload a document first.",
                    "document_modified": False
                }
        else:
            # Regular chat response
            chat_completion = client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": "You are a helpful assistant for resume editing. You can help users modify their resumes and provide guidance on resume writing best practices."
                    },
                    {
                        "role": "user",
                        "content": chat.message,
                    }
                ],
                model="llama-3.3-70b-versatile",
            )
            return {
                "response": chat_completion.choices[0].message.content,
                "document_modified": False
            }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat error: {e}")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    global CURRENT_DOCUMENT
    original_filename = file.filename
    name, ext = os.path.splitext(original_filename)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    saved_filename = f"{name}-{timestamp}{ext}"
    file_location = os.path.join(UPLOAD_FOLDER, saved_filename)

    with open(file_location, "wb") as f:
        content = await file.read()
        f.write(content)

    # If uploaded file is a PDF, also convert to DOCX and save
    # Only allow .pdf or .docx uploads
    if ext.lower() not in [".pdf", ".docx"]:
        raise HTTPException(status_code=400, detail="Only .pdf or .docx files are allowed")

    if ext.lower() == ".pdf":
        try:
            base_name = os.path.splitext(saved_filename)[0]
            docx_filename = f"{base_name}.docx"
            docx_location = os.path.join(UPLOAD_FOLDER, docx_filename)

            # Offload blocking conversion and apply a timeout guard
            await asyncio.wait_for(_convert_pdf_to_docx(file_location, docx_location), timeout=120)

            # Remove source PDF so only final .docx remains
            try:
                os.remove(file_location)
            except Exception:
                pass

            # Set as current document
            CURRENT_DOCUMENT = docx_filename
            
            return {"message": "File uploaded and converted to .docx successfully", "saved_filename": docx_filename}
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Conversion timed out")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Conversion failed: {e}")

    # If .docx, set as current document
    if ext.lower() == ".docx":
        CURRENT_DOCUMENT = saved_filename
        return {"message": "File uploaded successfully", "saved_filename": saved_filename}

@app.get("/preview/{filename}")
async def preview_file(filename: str):
    file_path = os.path.join(UPLOAD_FOLDER, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    _, ext = os.path.splitext(filename)
    ext = ext.lower()

    try:
        if ext == ".pdf":
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return {"filename": filename, "type": "pdf", "content": text[:10000]}
        elif ext == ".docx":
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            content = "\n".join(paragraphs)
            # Fallback: also try tables
            if not content.strip() and getattr(doc, "tables", None):
                rows = []
                for table in doc.tables:
                    for row in table.rows:
                        cells_text = [cell.text for cell in row.cells]
                        rows.append("\t".join(cells_text))
                content = "\n".join(rows)
            return {"filename": filename, "type": "docx", "content": content[:10000]}
        else:
            raise HTTPException(status_code=400, detail="Preview supports only .pdf or .docx")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not read file: {e}")

@app.get("/current-document")
async def get_current_document():
    """Get the currently active document filename"""
    return {"current_document": CURRENT_DOCUMENT}

@app.get("/current-preview")
async def get_current_document_preview():
    """Get the preview content of the currently active document"""
    if not CURRENT_DOCUMENT:
        raise HTTPException(status_code=404, detail="No active document found")
    
    file_path = os.path.join(UPLOAD_FOLDER, CURRENT_DOCUMENT)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Current document not found")
    
    _, ext = os.path.splitext(CURRENT_DOCUMENT)
    ext = ext.lower()
    
    try:
        if ext == ".pdf":
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return {
                "filename": CURRENT_DOCUMENT, 
                "type": "pdf", 
                "content": text[:10000],
                "is_current": True
            }
        elif ext == ".docx":
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            content = "\n".join(paragraphs)
            
            # Also extract table content
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    cells_text = [cell.text for cell in row.cells]
                    table_content.append(" | ".join(cells_text))
            
            if table_content:
                content += "\n\nTables:\n" + "\n".join(table_content)
            
            return {
                "filename": CURRENT_DOCUMENT, 
                "type": "docx", 
                "content": content[:10000],
                "is_current": True
            }
        else:
            raise HTTPException(status_code=400, detail="Preview supports only .pdf or .docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not read current document: {e}")

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found/") 
    return FileResponse(path=file_path, filename=filename, media_type="application/octet-stream")
