#!/usr/bin/env python3
"""
Test script for document modification functionality
"""

import os
import sys
from docx import Document

def create_test_document():
    """Create a test .docx document for testing"""
    doc = Document()
    
    # Add some test content
    doc.add_heading("Test Resume", 0)
    doc.add_paragraph("This is a test resume for testing the modification functionality.")
    doc.add_paragraph("Name: John Doe")
    doc.add_paragraph("Email: john.doe@example.com")
    doc.add_paragraph("Phone: (555) 123-4567")
    
    # Add a table
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Skill"
    table.rows[0].cells[1].text = "Level"
    table.rows[0].cells[2].text = "Years"
    table.rows[1].cells[0].text = "Python"
    table.rows[1].cells[1].text = "Expert"
    table.rows[1].cells[2].text = "5"
    
    # Save the document
    test_file = "test_resume.docx"
    doc.save(test_file)
    print(f"Created test document: {test_file}")
    return test_file

def test_document_modification():
    """Test the document modification function"""
    try:
        # Import the function directly
        import importlib.util
        spec = importlib.util.spec_from_file_location("main", "main.py")
        main_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(main_module)
        
        _modify_document = main_module._modify_document
        
        # Create a mock Groq client that matches the exact structure
        class MockGroqClient:
            def __init__(self):
                pass
            
            def chat(self):
                return self
            
            def completions(self):
                return self
            
            def create(self, **kwargs):
                class MockResponse:
                    def __init__(self):
                        self.choices = [type('obj', (object,), {'message': type('obj', (object,), {'content': '''{
                            "paragraphs": [
                                "Modified Test Resume",
                                "This resume has been updated according to your request.",
                                "Name: John Doe (Updated)",
                                "Email: john.doe.updated@example.com",
                                "Phone: (555) 987-6543"
                            ],
                            "tables": [
                                {
                                    "rows": [
                                        ["Skill", "Level", "Years"],
                                        ["Python", "Expert", "5"],
                                        ["JavaScript", "Intermediate", "3"]
                                    ]
                                }
                            ]
                        }'''})})]
                return MockResponse()
        
        # Test the modification
        test_file = create_test_document()
        
        try:
            client = MockGroqClient()
            success = _modify_document(test_file, "Update the resume with more professional language and add JavaScript skills", client)
            
            if success:
                print("✅ Document modification successful!")
                
                # Read and display the modified content
                doc = Document(test_file)
                print("\nModified document content:")
                for para in doc.paragraphs:
                    if para.text.strip():
                        print(f"- {para.text}")
                
                print("\nModified tables:")
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        print(f"  {' | '.join(row_text)}")
            else:
                print("❌ Document modification failed!")
                
        except Exception as e:
            print(f"❌ Error during testing: {e}")
            import traceback
            traceback.print_exc()
        
        finally:
            # Clean up
            if os.path.exists(test_file):
                os.remove(test_file)
                print(f"\nCleaned up test file: {test_file}")
                
    except Exception as e:
        print(f"❌ Error importing main module: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("Testing document modification functionality...")
    test_document_modification()
